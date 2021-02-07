import os, random
from docx import Document
import docx
import random
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt

if os.path.exists("C:\exam\demo.docx"):
    os.remove("C:\exam\demo.docx")
if os.path.exists("C:\exam\מבחן מעורבל.docx"):
    os.remove("C:\exam\מבחן מעורבל.docx")
arr_txt = [x for x in os.listdir("C:\exam") if x.endswith(".docx")]
print(int(len(arr_txt)/2))
questions = int(len(arr_txt)/2)
num_of_questions=input("כמה שאלות יהיו במבחן ?")
while int(num_of_questions) > questions:
    print('אין מספיק שאלות!!')
    num_of_questions = input("כמה שאלות יהיו במבחן ?")
range_from=range(1,int(num_of_questions)+1)
random_list=random.sample(range_from,int(num_of_questions))
Document().save('C:\exam\demo.docx')
master = Document('C:\exam\demo.docx')
composer = Composer(master)
for i in range(len(random_list)):
    paragraph=master.add_paragraph()
    run=paragraph.add_run("שאלה מספר"+" "+str(i+1)+" ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run.font.size = Pt(18)
    run.font.underline= True
    #font.name = 'Calibri'
    print(random_list[i])
    doc1 = Document("C:\exam\\"+str(random_list[i])+"_H_question.docx")
    composer.append(doc1)
    master.add_paragraph('')
    paragraph=master.add_paragraph('__________________________________________________________________________________________')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    master.add_paragraph('')
master.add_page_break()
paragraph=master.add_paragraph()
run=paragraph.add_run("מחוון")
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.font.size = Pt(24)
for i in range(len(random_list)):
    paragraph=master.add_paragraph()
    run=paragraph.add_run("תשובה לשאלה"+" "+str(i+1)+" ")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run.font.size = Pt(18)
    run.font.underline= True
    #font.name = 'Calibri'
    print(random_list[i])
    doc1 = Document("C:\exam\\"+str(random_list[i])+"_H_solution.docx")
    composer.append(doc1)
    master.add_paragraph('')
    paragraph=master.add_paragraph('__________________________________________________________________________________________')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    master.add_paragraph('')
print(random_list)
composer.save("C:\exam\מבחן מעורבל.docx")



